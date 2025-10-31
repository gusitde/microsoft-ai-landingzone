import os
import json
import docx
from azure.identity import DefaultAzureCredential
from azure.mgmt.resource import ResourceManagementClient

# Initialize Azure credentials and resource management client
credential = DefaultAzureCredential()
subscription_id = os.environ["AZURE_SUBSCRIPTION_ID"]
resource_client = ResourceManagementClient(credential, subscription_id)

def get_resources():
  resources = []
  for item in resource_client.resources.list():
    resources.append({
      'name': item.name,
      'type': item.type,
      'location': item.location,
      'resource_group': item.resource_group,
      'id': item.id
    })
  return resources

def create_docx(resources):
  doc = docx.Document()
  doc.add_heading('Azure Infrastructure As-Built Report', level=1)

  for resource in resources:
    doc.add_heading(resource['name'], level=2)
    doc.add_paragraph(f"Type: {resource['type']}")
    doc.add_paragraph(f"Location: {resource['location']}")
    doc.add_paragraph(f"Resource Group: {resource['resource_group']}")
    doc.add_paragraph(f"ID: {resource['id']}")
    doc.add_paragraph("\n")

  doc.save('as_built_report.docx')

def main():
  resources = get_resources()
  create_docx(resources)
  print("As-Built report generated successfully.")

if __name__ == "__main__":
  main()
